import re
from docx import Document


# Helper for formatting hymn numbers
def format_hymn_number(hymn_num):
    if len(hymn_num) in [2, 3]:
        return f"UMH {hymn_num}"
    elif len(hymn_num) == 4:
        if hymn_num.startswith('2'):
            return f"TFWS {hymn_num}"
        elif hymn_num.startswith('3'):
            return "see screens"
    return hymn_num

def parse_source_doc(file_path):
    """
    Parses the source Word document to extract service details, readings, and hymns.
    """
    data = {}
    doc = Document(file_path)

    # Initialize defaults
    data['is_communion_sunday'] = False

    # Helper to find text after a label
    # This is a basic implementation; we might need to iterate through paragraphs
    # and find the one *after* a bold label, or on the same line.

    full_text = "\n".join([p.text for p in doc.paragraphs])

    # Regex patterns for extraction (based on schema and PDF analysis)
    # Note: nuances like "16th Sunday" might be dynamic.
    # We look for patterns.

    # Service Details
    # Date: "September 28, 2025" -> Look for Month Day, Year pattern
    date_match = re.search(r'([A-Z][a-z]+ \d{1,2}, \d{4})', full_text)
    if date_match:
        data['date'] = date_match.group(1)

    # Service Time: "10:30 am"
    time_match = re.search(r'(\d{1,2}:\d{2} [ap]m)', full_text)
    if time_match:
        data['service_time'] = time_match.group(1)

    # Sunday Title: Often the first line or near "Sunday after Pentecost"
    # match anything line containing "Sunday"
    for p in doc.paragraphs:
        if "Sunday" in p.text and len(p.text) < 100: # heuristic
            raw_text = p.text

            # Stop at " +" if present (User request: stop at " +" or line break)
            if " +" in raw_text:
                raw_text = raw_text.split(" +")[0]

            # Remove Date: Month Day, Year
            raw_text = re.sub(r'[A-Z][a-z]+ \d{1,2}, \d{4}', '', raw_text)
            # Remove Time: HH:MM am/pm
            raw_text = re.sub(r'\d{1,2}:\d{2} [ap]m', '', raw_text)

            # Clean up extra whitespace and potential separators
            cleaned_title = re.sub(r'\s+', ' ', raw_text).strip(' -–—')

            data['sunday_title'] = cleaned_title
            break

    # Special/Series Title
    # Look for "Worship Series" and grab the next line
    for i, p in enumerate(doc.paragraphs):
        if "Worship Series" in p.text:
            # Look ahead for the title
            if i + 1 < len(doc.paragraphs):
                data['special_title'] = doc.paragraphs[i+1].text.strip('“"” ')
                break

    # Hymns (Extracting numbers and titles)
    # Pattern seen: Hymn—62, “All Creatures...”
    # We regex the whole text for ease, or we can iterate pars if order matters.
    # Let's iterate pars to keep them ordered 1, 2, 3...

    hymn_count = 0
    # Updated Regex to capture instructions after the title
    # Support "Hymn-123" OR "Communion Hymn-123" OR "Communion-123"
    # Group 1: Num, Group 2: Title, Group 3: Instructions (Optional)
    hymn_pattern = re.compile(r'(?:Hymn|Communion).{0,10}[—\-–](\d+).*?[“"”](.+?)[“"”]\s*(.*)$', re.IGNORECASE)

    # Doxology Pattern
    doxology_pattern = re.compile(r'\b(94|95)\b')

    # Flag to track if we just saw a "Communion" related header
    expecting_communion_hymn = False

    for p in doc.paragraphs:
        text = p.text.strip()

        # Check for Doxology
        dox_match = doxology_pattern.search(text)
        if dox_match and "Hymn" not in text and "Reading" not in text:
             data['doxology_num'] = format_hymn_number(dox_match.group(1))
             continue

        # Check for Context Headers BEFORE checking for Hymns
        # If we see "Sharing the Bread" or "Communion", set the flag for the NEXT hymn
        # We check this first so we don't accidentally match the instruction line itself as a hymn if it's not
        # Crucial Fix: Only look for context if we haven't found the communion hymn yet!
        # This prevents "Prayer after Communion" from incorrectly flagging the *next* hymn (e.g. Closing Hymn) as a communion hymn.
        if 'communion_hymn_num' not in data:
             if "sharing the bread" in text.lower() or "communion" in text.lower() or "holy mystery" in text.lower():
                expecting_communion_hymn = True

        match = hymn_pattern.search(text)
        if match:
            hymn_num = format_hymn_number(match.group(1))
            hymn_title = match.group(2)
            hymn_instr = match.group(3).strip("()")

            # Special case: Communion Hymn
            # 1. Check intent via keywords in THIS line
            # 2. Check intent via PREVIOUS context flag
            is_communion_intent = (
                "communion" in text.lower() or
                "table" in text.lower() or
                "bread" in hymn_title.lower() or
                expecting_communion_hymn
            )

            if is_communion_intent:
                 data['communion_hymn_num'] = hymn_num
                 data['communion_hymn_title'] = hymn_title
                 data['is_communion_sunday'] = True
                 # Reset flag so we don't capture the next hymn as communion too
                 expecting_communion_hymn = False
            else:
                hymn_count += 1
                data[f'hymn_{hymn_count}_num'] = hymn_num
                data[f'hymn_{hymn_count}_title'] = hymn_title
                data[f'hymn_{hymn_count}_instr'] = hymn_instr

    # Readings
    # Handle sequentially like hymns.
    # Pattern: "First Reading—...", "Gospel Reading—...", "Reading—..." (Case insensitive)

    reading_count = 0
    # Capture any line starting broadly with "Something Reading" or just "Reading"
    # Group 1: Verse, Group 2 (Optional): Translation parens
    reading_pattern = re.compile(r'(?:First|Second|Gospel|Hebrew Bible)?\s*Reading[—\-–](.+?)(?:\s+\((.+?)\))?$', re.IGNORECASE)

    for p in doc.paragraphs:
        text = p.text.strip()
        match = reading_pattern.search(text)
        if match:
            reading_count += 1

            r_verse = match.group(1).strip()
            # Strip parens from translation if present
            # Group 2 might be None now if not captured
            raw_trans = match.group(2)
            r_trans = raw_trans.strip("()") if raw_trans else ""

            # Assign to reading_1, reading_2...
            data[f'reading_{reading_count}_verse'] = r_verse
            data[f'reading_{reading_count}_translation'] = r_trans

    # Final check: If we found a communion hymn num, ensure the flag is true
    if 'communion_hymn_num' in data:
        data['is_communion_sunday'] = True

    return data

def parse_email_text(text, source_type="organist"):
    data = {}
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    # Map old generic types if passed
    if source_type == "music_1": source_type = "organist"
    if source_type == "music_2": source_type = "choir"

    if source_type == "organist":
        keywords = {
            'Prelude': 'prelude',
            'New Spirit': 'new_spirit',
            'Communion': 'communion_piece',
            'Postlude': 'postlude',
            'Exiting': 'exit_music',
            'Exit Music': 'exit_music',
            'Exit': 'exit_music'
        }
    else: # choir
        keywords = {
            'Introit': 'introit',
            'Anthem': 'anthem',
            'Prayer Response': 'prayer_response',
            'Benediction Response': 'benediction_response', # Added
            'Benediction': 'benediction_response',
            'Choral Benediction Response': 'benediction_response',
            'Choral Benediction': 'benediction_response',
            'Response': 'prayer_response' # Fallback
        }

    for i, line in enumerate(lines):
        detected_key = None
        detected_prefix = None

        clean_line = line.strip().lower()

        for key, prefix in keywords.items():
            k_low = key.lower()
            # Strict Header Matching:
            # 1. Exact match "Prelude"
            # 2. Starts with "Prelude:" or "Prelude-"
            is_header = False
            if clean_line == k_low:
                is_header = True
            elif clean_line.startswith(k_low + ":") or clean_line.startswith(k_low + "-"):
                is_header = True

            if is_header:
                detected_key = key
                detected_prefix = prefix
                break

        if detected_prefix:
            content_lines = []
            for j in range(i+1, len(lines)):
                next_line = lines[j]
                next_clean = next_line.strip().lower()

                # Check if next line is a keyword (using strict logic too!)
                is_next_keyword = False
                for k in keywords:
                    k_low = k.lower()
                    if next_clean == k_low:
                        is_next_keyword = True
                        break
                    elif next_clean.startswith(k_low + ":") or next_clean.startswith(k_low + "-"):
                        is_next_keyword = True
                        break

                if is_next_keyword:
                    break
                content_lines.append(next_line)

            if not content_lines:
                continue

            # Analyze content lines
            # Strategy:
            # 1. First line usually contains Title + Composer (separated by tabs/spaces)
            # 2. Subsequent lines are checked:
            #    - If they start with '[' or '(', they are IGNORED (per user request).
            #    - Otherwise, they are treated as Personnel.

            main_line = content_lines[0]
            extra_lines = content_lines[1:] if len(content_lines) > 1 else []

            # Filter extra lines
            personnel_lines = []
            for l in extra_lines:
                clean_l = l.strip()
                # Ignore lines starting with brackets or parentheses
                if clean_l.startswith('[') or clean_l.startswith('('):
                    continue
                # Otherwise, it's personnel
                personnel_lines.append(clean_l)

            # Clean items in main line (stripping brackets if any)
            main_line = re.sub(r'[\(\[].*?[\)\]]', '', main_line).strip()

            # Personnel string
            personnel_text = "; ".join(personnel_lines)

            title = ""
            composer = ""

            # Priority 1: " - "
            if " - " in main_line:
                parts = main_line.split(" - ")
                title = parts[0].strip()
                composer = parts[1].strip() if len(parts) > 1 else ""

            # Priority 2: Tabs or Multiple Spaces (Choir Email format)
            elif re.search(r'\t|\s{2,}', main_line):
                parts = re.split(r'\t|\s{2,}', main_line)
                parts = [p.strip() for p in parts if p.strip()]
                title = parts[0] if parts else ""
                composer = parts[1] if len(parts) > 1 else ""

            # Priority 3: Parenthesis ending title, followed by text (Author/Composer)
            # e.g. "God’s Bright Star (Hodie) Piae Cantiones" -> Title: "God’s Bright Star (Hodie)", Comp: "Piae Cantiones"
            elif re.search(r'\)\s+[A-Z]', main_line):
                # Split at the last closing parenthesis that is followed by a space and capital letter
                # We use a regex to find the split point
                parts = re.split(r'(?<=\))\s+(?=[A-Z])', main_line, 1)
                title = parts[0].strip()
                composer = parts[1].strip() if len(parts) > 1 else ""

            # Priority 4: " by "
            elif " by " in main_line:
                parts = main_line.split(" by ")
                title = parts[0].strip()
                composer = parts[1].strip() if len(parts) > 1 else ""

            # Priority 5: ", harm." or ", arr." or ", ed."
            # e.g. "Tune Name, harm. Bach" -> Title: "Tune Name", Comp: "harm. Bach"
            elif re.search(r', (harm|arr|ed)\.', main_line):
                # Split at the first occurrence of these keywords
                # We want to keep the keyword in the composer part, so we split on the comma
                match = re.search(r', (harm|arr|ed)\.', main_line)
                if match:
                    split_idx = match.start()
                    title = main_line[:split_idx].strip()
                    composer = main_line[split_idx+1:].strip() # Skip comma

            else:
                # If single line, it's just the Title.
                title = main_line
                # If there were extra lines, maybe the FIRST extra line was actually the composer?
                # User said: "header, line break, title, spaces, composer, line break, further details"
                # But also: "header, line break, title, line break, composer" logic from before?
                # Let's stick to the "whitespace on line 1" rule as primary.
                # If line 1 has no composer, and we have extra lines, let's NOT assume line 2 is composer
                # unless we want to fallback.
                # User's specific request: "further composer information... handle as separate variable"
                # So assume Line 1 has the core info. Line 2+ is details.
                composer = ""

            # Sanitize
            if title.lower() == "none": title = ""
            if composer.lower() == "none": composer = ""
            if personnel_text.lower() == "none": personnel_text = ""

            data[f'{detected_prefix}_title'] = title or ""
            data[f'{detected_prefix}_composer'] = composer or ""
            data[f'{detected_prefix}_personnel'] = personnel_text or ""

    return data
